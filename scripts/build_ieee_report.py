#!/usr/bin/env python3
"""Build the editable IEEE-style project report and its figures."""

from __future__ import annotations

import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "output" / "docx"
OUT_PDF = ROOT / "output" / "pdf"
ASSETS = ROOT / "output" / "assets"
REPORT_PATH = OUT_DOCX / "X24244066_Clickstream_Analytics_Report.docx"
ARCH_PATH = ASSETS / "clickstream_architecture.png"
PERF_PATH = ASSETS / "emr_performance_report.png"
DASHBOARD_PATH = ASSETS / "dashboard_student.png"

INK = RGBColor(0x12, 0x18, 0x26)
MUTED = RGBColor(0x4B, 0x55, 0x63)
BLUE = RGBColor(0x1F, 0x5F, 0xB8)
YELLOW = RGBColor(0xFF, 0xEB, 0x9C)


def font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline, label, sublabel=None, accent=None):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    if accent:
        draw.rounded_rectangle((x1, y1, x1 + 14, y2), radius=7, fill=accent)
    f1 = font(26, True)
    f2 = font(20, False)
    center = (x1 + x2) / 2
    bbox = draw.textbbox((0, 0), label, font=f1)
    draw.text((center - (bbox[2] - bbox[0]) / 2, y1 + 19), label, font=f1, fill="#111827")
    if sublabel:
        lines = sublabel if isinstance(sublabel, list) else [sublabel]
        y = y1 + 57
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=f2)
            draw.text((center - (bbox[2] - bbox[0]) / 2, y), line, font=f2, fill="#4B5563")
            y += 26


def arrow(draw, start, end, label=None):
    draw.line((start, end), fill="#475569", width=5)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 16
    for delta in (2.6, -2.6):
        tip = (end[0] + size * math.cos(angle + delta), end[1] + size * math.sin(angle + delta))
        draw.line((end, tip), fill="#475569", width=5)
    if label:
        f = font(18, False)
        mx = (start[0] + end[0]) / 2
        my = (start[1] + end[1]) / 2 - 24
        bbox = draw.textbbox((0, 0), label, font=f)
        draw.rounded_rectangle((mx - (bbox[2] - bbox[0]) / 2 - 7, my - 2,
                                mx + (bbox[2] - bbox[0]) / 2 + 7, my + 23), radius=5, fill="white")
        draw.text((mx - (bbox[2] - bbox[0]) / 2, my), label, font=f, fill="#475569")


def build_architecture():
    img = Image.new("RGB", (1050, 1220), "white")
    d = ImageDraw.Draw(img)
    title = "Implemented Lambda Architecture"
    tf = font(34, True)
    tb = d.textbbox((0, 0), title, font=tf)
    d.text(((1050 - (tb[2] - tb[0])) / 2, 25), title, font=tf, fill="#111827")

    rounded_box(d, (70, 100, 440, 205), "#EFF6FF", "#2563EB", "REES46 dataset", "Public electronics events", "#2563EB")
    rounded_box(d, (610, 100, 980, 205), "#F5F3FF", "#7C3AED", "Python producer", ["Schema validation", "Controlled replay clock"], "#7C3AED")
    arrow(d, (440, 152), (610, 152), "CSV.GZ")

    rounded_box(d, (330, 285, 720, 395), "#FFF7ED", "#EA580C", "Amazon Kinesis", ["Session-based partition key", "Live ingestion stream"], "#EA580C")
    arrow(d, (795, 205), (620, 285), "canonical JSON")

    d.rounded_rectangle((35, 445, 1015, 1000), radius=25, fill="#FAFAFA", outline="#CBD5E1", width=3)
    d.text((65, 463), "Elastic processing and storage boundary", font=font(24, True), fill="#334155")

    rounded_box(d, (70, 535, 450, 650), "#ECFDF5", "#059669", "AWS Lambda", ["Minute-window counters", "Latency health metric"], "#059669")
    rounded_box(d, (600, 535, 980, 650), "#EFF6FF", "#2563EB", "Amazon S3 raw", ["Historical canonical data", "Five benchmark copies"], "#2563EB")
    arrow(d, (430, 395), (270, 535), "event batches")
    arrow(d, (620, 395), (790, 535), "history upload")

    rounded_box(d, (70, 735, 450, 850), "#F0FDFA", "#0F766E", "Amazon DynamoDB", ["One-minute product buckets", "48-hour TTL"], "#0F766E")
    rounded_box(d, (600, 735, 980, 850), "#FDF2F8", "#BE185D", "Amazon EMR + Spark", ["Five-minute baselines", "1 vs. 8 partitions"], "#BE185D")
    arrow(d, (260, 650), (260, 735))
    arrow(d, (790, 650), (790, 735))

    rounded_box(d, (600, 895, 980, 975), "#FFF7ED", "#C2410C", "S3 analytics", "Parquet + JSON baselines", "#C2410C")
    arrow(d, (790, 850), (790, 895))

    rounded_box(d, (250, 1050, 800, 1165), "#F8FAFC", "#334155", "Python serving merge", ["Current speed view + historical baseline", "Trending products and funnel drop-off"], "#334155")
    arrow(d, (260, 850), (430, 1050), "recent metrics")
    arrow(d, (790, 975), (650, 1050), "baseline")

    img.save(ARCH_PATH, quality=95)


def build_performance_figure():
    img = Image.new("RGB", (1050, 760), "white")
    d = ImageDraw.Draw(img)
    title = "PySpark Batch Benchmark (999,825 Events)"
    ftitle = font(32, True)
    bb = d.textbbox((0, 0), title, font=ftitle)
    d.text(((1050 - (bb[2] - bb[0])) / 2, 25), title, font=ftitle, fill="#111827")
    d.text((78, 88), "Same EMR cluster | two initial workers | 290 MB input", font=font(23), fill="#4B5563")

    base_y = 545
    axis_x = 130
    axis_w = 790
    d.line((axis_x, 165, axis_x, base_y), fill="#64748B", width=3)
    d.line((axis_x, base_y, axis_x + axis_w, base_y), fill="#64748B", width=3)
    for sec in range(0, 81, 20):
        y = base_y - int(sec / 80 * 340)
        d.line((axis_x, y, axis_x + axis_w, y), fill="#E2E8F0", width=2)
        label = str(sec)
        b = d.textbbox((0, 0), label, font=font(21))
        d.text((axis_x - 18 - (b[2] - b[0]), y - 12), label, font=font(21), fill="#475569")
    d.text((20, 310), "seconds", font=font(22, True), fill="#475569")

    bars = [(315, 70, "1 partition", "Sequential", "#6B7280"), (650, 58, "8 partitions", "Parallel", "#2563EB")]
    for x, value, label1, label2, color in bars:
        h = int(value / 80 * 340)
        d.rounded_rectangle((x - 105, base_y - h, x + 105, base_y), radius=14, fill=color)
        val = f"{value} s"
        b = d.textbbox((0, 0), val, font=font(30, True))
        d.text((x - (b[2] - b[0]) / 2, base_y - h - 48), val, font=font(30, True), fill="#111827")
        for idx, line in enumerate((label1, label2)):
            b = d.textbbox((0, 0), line, font=font(23, idx == 0))
            d.text((x - (b[2] - b[0]) / 2, base_y + 20 + idx * 31), line, font=font(23, idx == 0), fill="#334155")

    d.rounded_rectangle((130, 650, 920, 725), radius=18, fill="#EFF6FF", outline="#93C5FD", width=2)
    note = "Observed result: 12 s faster | 17.1% lower duration | speedup = 1.207x"
    b = d.textbbox((0, 0), note, font=font(23, True))
    d.text(((1050 - (b[2] - b[0])) / 2, 673), note, font=font(23, True), fill="#1E3A8A")
    img.save(PERF_PATH, quality=95)


def set_run(run, *, size=9.5, bold=False, italic=False, color=INK, font_name="Times New Roman"):
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_run(run, fill="FFEB9C"):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    run._element.get_or_add_rPr().append(shd)


def body(doc, text, *, first=False, after=2.2, italic=False, bold_lead=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.first_line_indent = Inches(0 if first else 0.16)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run(r2, italic=italic)
    else:
        r = p.add_run(text)
        set_run(r, italic=italic)
    return p


def equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run(r, size=9.2, italic=True)
    return p


def heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text.upper())
    set_run(r, size=10, bold=True)
    return p


def subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, size=9.5, bold=True, italic=True)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, size=8, italic=False)
    return p


def add_figure(doc, path, width, cap):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(str(path), width=Inches(width))
    caption(doc, cap)


def add_cropped_figure(doc, path, width, height, cap, *, left, top, right, bottom):
    """Insert a picture cropped to a percentage-based source rectangle."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    shape = p.add_run().add_picture(str(path), width=Inches(width), height=Inches(height))
    blip_fill = shape._inline.graphic.graphicData.pic.blipFill
    src_rect = OxmlElement("a:srcRect")
    src_rect.set("l", str(int(left * 1000)))
    src_rect.set("t", str(int(top * 1000)))
    src_rect.set("r", str(int(right * 1000)))
    src_rect.set("b", str(int(bottom * 1000)))
    blip_fill.insert(1, src_rect)
    caption(doc, cap)


def cell_margins(cell, top=55, start=70, bottom=55, end=70):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_borders(table, color="AAB2BD", size="4"):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)
        borders.append(el)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_VERTICAL.CENTER
    table.autofit = False
    table_borders(table)
    total = sum(widths)
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(total))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width))
        grid.append(gc)
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i] / 1440)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E8EEF5")
        cell._tc.get_or_add_tcPr().append(shd)
        cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(text))
        set_run(r, size=7.6, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i] / 1440)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cells[i])
            tcW = cells[i]._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(widths[i]))
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run(r, size=7.5)
        trPr = table.rows[-1]._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height = Pt(18)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def set_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.66)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.28)


def set_columns(section, count=2, space=360):
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")
    if cols:
        cols_el = cols[0]
    else:
        cols_el = OxmlElement("w:cols")
        sectPr.append(cols_el)
    cols_el.set(qn("w:num"), str(count))
    cols_el.set(qn("w:space"), str(space))
    cols_el.set(qn("w:equalWidth"), "1")


def page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = True
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)
    set_run(run, size=8, color=MUTED)


def build_report():
    doc = Document()
    first = doc.sections[0]
    set_page(first)
    set_columns(first, 1)
    page_number(first)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.2)
    normal.paragraph_format.line_spacing = 1.0

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(7)
    r = title.add_run("Scalable Real-Time E-Commerce\nClickstream Analytics")
    set_run(r, size=18, bold=True)

    auth = doc.add_paragraph()
    auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth.paragraph_format.space_after = Pt(1)
    r = auth.add_run("SHARMILA RAMARAJ  |  X24244066")
    set_run(r, size=10.5)

    partner = doc.add_paragraph()
    partner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    partner.paragraph_format.space_after = Pt(3)
    r = partner.add_run("PARTNER: [INSERT NAME AND STUDENT ID]")
    set_run(r, size=10, bold=True)
    shade_run(r)

    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org.paragraph_format.space_after = Pt(8)
    r = org.add_run("MSc Cloud Computing, National College of Ireland | Scalable Cloud Programming | August 2026")
    set_run(r, size=9, italic=True, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Abstract-")
    set_run(r, size=9, bold=True, italic=True)
    abstract = (
        " E-commerce teams need historical context and current behaviour at the same time. This project implements a reproducible AWS Lambda architecture that replays a public REES46 electronics clickstream into Amazon Kinesis, processes recent events with AWS Lambda, stores minute aggregates in DynamoDB, and computes historical five-minute product baselines with PySpark on Amazon EMR. A Python serving layer and public AWS-hosted dashboard merge both views to identify unusually trending products and quantify funnel drop-off. From 200,000 source rows, 199,965 valid events were normalised; five S3 copies produced a 999,825-event, approximately 290 MB benchmark. On the same two-worker EMR cluster, the one-partition reference completed in 70 s and the eight-partition run in 58 s, a 1.207x speedup and 17.1% lower duration. A deterministic live test verified Kinesis-to-Lambda processing and produced the expected DynamoDB window rows. The results demonstrate the analytical value of combining recent activity with historical norms while showing that fixed Spark and storage overheads limit speedup on modest inputs."
    )
    r = p.add_run(abstract)
    set_run(r, size=9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Keywords-")
    set_run(r, size=9, bold=True, italic=True)
    r = p.add_run(" clickstream analytics, Lambda architecture, Amazon Kinesis, AWS Lambda, DynamoDB, PySpark, Amazon EMR, auto-scaling")
    set_run(r, size=9)

    second = doc.add_section(WD_SECTION.CONTINUOUS)
    set_page(second)
    set_columns(second, 2, 360)
    page_number(second)

    heading(doc, "I. Introduction")
    body(doc, "Online stores generate product views, cart actions, removals and purchases continuously. Historical reports can show long-term popularity but may hide a sudden spike, while a live counter can mistake normal popularity for an anomaly. The practical question addressed here is: Which products have unusually high activity during the latest five-minute window, and is the larger current funnel drop-off view-to-cart or cart-to-purchase?", first=True)
    body(doc, "The resulting view supports marketing promotion, stock checking, user-experience investigation and platform operations. It is decision support only: it does not identify customers, change prices or contact users. The main contribution is an end-to-end Python implementation with reproducible input, separate batch and speed paths, an explicit serving merge, and measured parallel execution on AWS Academy infrastructure.")

    subheading(doc, "A. Objectives")
    body(doc, "The objectives were to: (1) build a controlled replay producer; (2) compute historical baselines with distributed PySpark; (3) calculate recent product and funnel metrics from Kinesis; (4) combine current and historical views; (5) configure elastic AWS processing; and (6) compare sequential-reference and parallel execution using the same input and cluster.", first=True)

    heading(doc, "II. Background and Design Rationale")
    body(doc, "The Lambda architecture separates a complete, recomputable batch view from a low-latency speed view and merges them for queries [1]. This separation fits the use case: batch computation supplies accurate historical norms, whereas recent counters reveal changes early. Spark provides distributed, fault-tolerant transformations derived from the resilient distributed dataset model [2]. Kinesis and Lambda provide a managed event path in which Lambda polls the stream and invokes the consumer for available records [3].", first=True)
    body(doc, "A single stream-only design would be simpler but would require durable state, replay and historical recomputation to answer the same question. A batch-only design would be accurate but stale. The selected hybrid therefore prioritises both correctness and freshness while remaining small enough for a Learner Lab.")

    heading(doc, "III. Use Case and Data")
    subheading(doc, "A. Operational Metrics")
    body(doc, "For product p in the latest five minutes, recent engagement is weighted so that actions nearer to conversion contribute more:", first=True)
    equation(doc, "trend_score(p) = views(p) + 3 carts(p) + 5 purchases(p).")
    body(doc, "The batch layer computes the historical mean of the same score over observed five-minute windows. Activity lift is the current score divided by this historical mean; a product is labelled unusually trending when lift is at least 2.0. This rule detects relative change, so a normally popular product is not automatically treated as anomalous.", first=True)
    equation(doc, "view-to-cart drop-off = 1 - cart sessions / view sessions")
    equation(doc, "cart-to-purchase drop-off = 1 - purchase sessions / cart sessions.")
    body(doc, "These ratios are window indicators, not proof that a named customer permanently abandoned checkout. A separate cart signal is generated after 15 replay-clock minutes if a session has neither purchased nor removed that product.", first=True)

    subheading(doc, "B. Dataset Preparation")
    body(doc, "The REES46 source contains anonymised e-commerce behavioural events and provides a compressed electronics event file [4]. Its fields include event time and type, product/category, brand, price, user and session identifiers. The source contained 885,129 events. The experiment selected the first 200,000 source rows, retained 199,965 valid events, and rejected 35 rows (0.0175%) without a session identifier because session funnels cannot be attributed reliably without it.", first=True)
    add_table(doc, ["Item", "Verified value"], [
        ("Source events", "885,129"),
        ("Selected / valid", "200,000 / 199,965"),
        ("Rejected", "35 (0.0175%)"),
        ("Canonical JSONL", "58 MB"),
        ("Benchmark", "999,825 events; ~290 MB"),
    ], [2360, 2680])

    heading(doc, "IV. System Architecture")
    add_figure(doc, ARCH_PATH, 3.42, "Fig. 1. Implemented batch, speed and serving paths.")
    body(doc, "The producer normalises the historical CSV.GZ file into canonical JSON, preserves the source timestamp and assigns a replay-clock event time. Session ID is the Kinesis partition key, preserving per-session order within a shard. The live branch invokes a Python Lambda and uses atomic DynamoDB updates for product and health counters in one-minute buckets with time-to-live expiry. The historical branch stores canonical input in S3, where EMR Spark aggregates five-minute windows and writes Parquet plus consolidated JSON baselines.", first=True)
    body(doc, "The serving module reads the recent view and historical product rows, calculates lift, orders products by relative change, and returns a JSON analytical view. A responsive static dashboard on an Amazon S3 website presents the same verified values as product activity, a three-stage funnel, service status and batch-performance panels. CloudWatch and retained EMR logs provide operational evidence.")

    heading(doc, "V. Implementation")
    subheading(doc, "A. Reproducible Ingestion")
    body(doc, "The Python producer accepts CSV or compressed CSV, maps source event types to the canonical schema, validates timezone-aware timestamps and required identifiers, and supports a record limit, target replay rate and no-sleep benchmark mode. The 199,965-event normalisation completed in 16.804 s, equivalent to approximately 11,900 records/s for local transformation and file output. This figure is not presented as end-to-end Kinesis throughput.", first=True)

    subheading(doc, "B. Batch Layer")
    body(doc, "The PySpark job reads JSON from S3, filters invalid rows, groups product events into five-minute windows, pivots event types and calculates activity scores. A second aggregation produces mean views, carts, purchases and trend score per product, while distinct session counts provide historical funnel context. Outputs are written as compressed Parquet for efficient analytical access and JSON for the serving demonstration.", first=True)

    subheading(doc, "C. Speed and Serving Layers")
    body(doc, "The Kinesis event-source mapping invokes the Lambda in batches. For each supported event the function increments the corresponding DynamoDB attribute and a health row containing event count and accumulated processing latency. Partial batch failure identifiers are returned so failed Kinesis records can be retried. The deterministic eight-event fixture produced three rows: two product aggregates and one health row.", first=True)
    body(doc, "The local stateful implementation maintains an exact five-minute deque and session counters for repeatable tests; the AWS Lambda persists equivalent one-minute buckets. The serving merge combines the latest buckets with batch averages, allowing recent counts to be interpreted in context. The student-built dashboard uses straightforward tables, progress bars and labelled funnel stages so the verified values remain easy to interpret.")

    subheading(doc, "D. Elasticity and Fault Handling")
    body(doc, "The EMR 7.13.0 cluster used Spark 3.5.6, one primary node, one core worker and one task worker at launch. Managed scaling was configured for a minimum of two and maximum of four worker instances, with at most one core worker. AWS documents that managed scaling evaluates cluster workload and can increase or decrease core/task capacity [5]. Lambda concurrency can expand with incoming event batches within account and event-source limits. The cluster also used a 20-minute idle-termination control and was manually terminated after measurement.", first=True)

    heading(doc, "VI. Experimental Method")
    body(doc, "Correctness was checked first with eight deterministic events covering two products and four sessions. After the event-source mapping reached Enabled, the replay populated DynamoDB and the serving JSON was compared with the known expected funnel and trend values. Seven automated local tests covered models, producer behaviour, speed logic, sequential baseline and serving merge.", first=True)
    body(doc, "For batch performance, five S3-side copies of the 199,965-event canonical subset formed one 999,825-event input of approximately 290 MB. Both runs used the same EMR cluster and program. The reference forced one Spark partition; the parallel case repartitioned by product into eight partitions. Duration was taken from completed EMR step timings. Speedup was calculated as T(1 partition) / T(configuration). This controls the data and infrastructure but varies partitions rather than worker count.")

    heading(doc, "VII. Results")
    subheading(doc, "A. Real-Time Analytical Answer")
    body(doc, "In the controlled serving result, product 200 scored 5 against a historical five-minute mean of 2.5. Its activity lift was 2.0, so it was flagged as unusually trending. Product 100 scored 13 against a mean of 13; lift was 1.0 and it was not flagged. Across the four view sessions, three reached cart and one reached purchase. Therefore view-to-cart drop-off was 25%, while cart-to-purchase drop-off was 66.67%. The larger immediate loss was after cart, suggesting that product/checkout friction should be investigated before adding more top-of-funnel promotion.", first=True)

    subheading(doc, "B. Serving Dashboard")
    add_cropped_figure(
        doc,
        DASHBOARD_PATH,
        3.28,
        2.26,
        "Fig. 2. Public AWS dashboard with the corrected session funnel.",
        left=15.6,
        top=29.8,
        right=2.7,
        bottom=29.3,
    )
    body(doc, "The public S3 website places the real-time answer first: product 200 is shown at 2.0x activity lift, with its current score, baseline and event composition visible in a plain results table. The funnel keeps all three stages and both drop-off percentages readable at desktop and mobile widths. A numbered pipeline panel explains how records move from the replay producer through Kinesis and Lambda to DynamoDB.", first=True)

    subheading(doc, "C. Batch Performance")
    add_figure(doc, PERF_PATH, 3.42, "Fig. 3. Same-cluster partitioning comparison on EMR.")
    add_table(doc, ["Configuration", "Part.", "Time", "Speedup"], [
        ("Sequential reference", "1", "70 s", "1.000x"),
        ("Parallel execution", "8", "58 s", "1.207x"),
    ], [2200, 760, 920, 1160])
    body(doc, "The parallel run completed 12 s faster, reducing elapsed duration by 17.1%. Output verification found two Parquet part files, one consolidated JSON file, success markers and approximately 9.75 MB of results. The 1.207x gain is useful but well below linear because Spark startup, S3 I/O, JSON parsing, shuffle coordination, task scheduling and final coalescing remain fixed or partly serial. The modest input also makes distributed overhead a larger fraction of total time.", first=True)

    heading(doc, "VIII. Critical Analysis")
    body(doc, "The strongest aspect is semantic reproducibility: a public dataset, deterministic fixture, explicit formulas and identical benchmark input make the result explainable. The Lambda architecture answers a question that neither layer answers alone. The unusual-trend rule prevents raw popularity from dominating, and the funnel result provides an actionable location for investigation.", first=True)
    body(doc, "The experiment nevertheless has important limits. First, one versus eight partitions on the same two-worker cluster is a parallelism comparison, not a controlled node-scaling experiment; therefore worker efficiency is not claimed. Second, the managed-scaling policy was configured and verified, but the 58-70 s jobs did not maintain demand long enough to prove an observed scale-out. Third, the cloud correctness replay contained only eight events, while the larger dataset evidence applies to preprocessing and EMR batch execution rather than sustained Kinesis load. Fourth, DynamoDB atomic increments are not idempotent, so a successful record retried after a partial failure could be counted twice.")
    body(doc, "Production improvements would add an event ID and deduplication record, emit percentile latency rather than only accumulated latency, retain raw stream records through a managed delivery path, and run repeated trials across several worker counts and data sizes. A longer backlog-generating load test should capture actual scaling timestamps, Kinesis iterator age, Lambda errors/throttles, executor utilisation and cost. Structured Streaming is also a viable future consolidation path because it supports event-time windows, checkpointing and incremental execution [6], although it would change the deliberately separate Lambda speed layer used here.")

    heading(doc, "IX. Conclusion")
    body(doc, "The project implements a compact but genuine AWS clickstream analytics pipeline. Kinesis and Lambda supply fresh counters, S3 and EMR provide recomputable history, DynamoDB stores speed metrics, and a Python serving merge plus dashboard turns both into product-trend and funnel decisions. The verified demonstration identified product 200 as unusually trending and located the larger drop-off between cart and purchase. The EMR test showed a measured 1.207x speedup from partitioned execution, while the analysis explains why the gain is not linear and avoids treating configuration as proof of scaling. The next priority is a longer, repeated load experiment with idempotent stream updates and observed scale events.", first=True)

    heading(doc, "References")
    refs = [
        "[1] N. Marz and J. Warren, Big Data: Principles and Best Practices of Scalable Real-Time Data Systems. Manning, 2015.",
        "[2] M. Zaharia et al., 'Resilient Distributed Datasets: A Fault-Tolerant Abstraction for In-Memory Cluster Computing,' Proc. NSDI, pp. 15-28, 2012. usenix.org/conference/nsdi12/technical-sessions/presentation/zaharia",
        "[3] Amazon Web Services, 'Tutorial: Using Lambda with Kinesis Data Streams,' AWS Lambda Developer Guide. [Online]. Available: docs.aws.amazon.com/lambda/latest/dg/with-kinesis-example.html. [Accessed: Aug. 3, 2026].",
        "[4] REES46, 'Events in an electronics and home-appliance store,' REES46 Datasets. [Online]. Available: data.rees46.com. [Accessed: Aug. 3, 2026].",
        "[5] Amazon Web Services, 'Using managed scaling in Amazon EMR,' Amazon EMR Management Guide. [Online]. Available: docs.aws.amazon.com/emr/latest/ManagementGuide/emr-managed-scaling.html. [Accessed: Aug. 3, 2026].",
        "[6] Apache Software Foundation, 'Structured Streaming Programming Guide, Spark 3.5.6.' [Online]. Available: spark.apache.org/docs/3.5.6/structured-streaming-programming-guide.html. [Accessed: Aug. 3, 2026].",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.16)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(1.2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(ref)
        set_run(r, size=7.0)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run("Submission links- ")
    set_run(r, size=6.8, bold=True)
    r = p.add_run("Repository: github.com/Sharmila-Ramaraj/scalable-clickstream-analytics")
    set_run(r, size=6.6)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run("AWS dashboard: http://scalable-real-time-clickstream-analytics-x24244066.s3-website-us-east-1.amazonaws.com  |  ")
    set_run(r, size=6.0)
    r = p.add_run("Video: [INSERT YOUTUBE OR ONEDRIVE URL]")
    set_run(r, size=6.0)
    shade_run(r)

    core = doc.core_properties
    core.title = "Scalable Real-Time E-Commerce Clickstream Analytics"
    core.author = "Sharmila Ramaraj (X24244066)"
    core.subject = "Scalable Cloud Programming continuous assessment"
    core.keywords = "AWS, Kinesis, Lambda, DynamoDB, EMR, PySpark, clickstream"
    doc.save(REPORT_PATH)


def main():
    OUT_DOCX.mkdir(parents=True, exist_ok=True)
    OUT_PDF.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    build_architecture()
    build_performance_figure()
    build_report()
    print(REPORT_PATH)


if __name__ == "__main__":
    main()
